import importlib.util
import unittest
from io import BytesIO

from study_assistant.extraction import (
    extract_from_bytes,
    extract_from_plain_text,
    get_tesseract_command,
    is_tesseract_available,
    normalize_text,
)

DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None


class ExtractionTests(unittest.TestCase):
    def test_normalize_text_preserves_paragraph_breaks(self):
        text = " First   line \r\n\r\n\r\n Second\tline "

        self.assertEqual(normalize_text(text), "First line\n\nSecond line")

    def test_extract_from_plain_text_reports_character_count(self):
        result = extract_from_plain_text("Embeddings find related ideas.", "notes")

        self.assertTrue(result.has_text)
        self.assertEqual(result.source_type, "text")
        self.assertEqual(result.metadata["characters"], len(result.text))

    def test_extract_text_file_from_bytes(self):
        result = extract_from_bytes("study_notes.txt", b"OCR reads images.\nEmbeddings search text.")

        self.assertEqual(result.source_name, "study_notes.txt")
        self.assertIn("Embeddings search text.", result.text)
        self.assertEqual(result.warnings, [])

    def test_unsupported_file_type_returns_warning(self):
        result = extract_from_bytes("archive.zip", b"PK")

        self.assertFalse(result.has_text)
        self.assertEqual(result.source_type, "unsupported")
        self.assertIn("Unsupported file type", result.warnings[0])

    def test_empty_file_returns_warning(self):
        result = extract_from_bytes("notes.txt", b"")

        self.assertFalse(result.has_text)
        self.assertEqual(result.source_type, "empty")
        self.assertIn("empty", result.warnings[0].lower())

    def test_legacy_doc_directs_user_to_docx(self):
        result = extract_from_bytes("notes.doc", b"\xd0\xcf\x11\xe0legacy")

        self.assertFalse(result.has_text)
        self.assertEqual(result.source_type, "unsupported")
        self.assertIn(".docx", result.warnings[0])

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx is not installed")
    def test_extract_docx_reads_paragraphs_and_tables(self):
        import docx

        document = docx.Document()
        document.add_paragraph("Embeddings convert text into vectors.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "OCR"
        table.rows[0].cells[1].text = "Reads images"
        buffer = BytesIO()
        document.save(buffer)

        result = extract_from_bytes("study_notes.docx", buffer.getvalue())

        self.assertEqual(result.source_type, "docx")
        self.assertIn("Embeddings convert text into vectors.", result.text)
        self.assertIn("Reads images", result.text)
        self.assertEqual(result.metadata["tables"], 1)

    def test_tesseract_probe_returns_boolean(self):
        self.assertIsInstance(is_tesseract_available(), bool)

    def test_tesseract_command_probe_returns_string_or_none(self):
        command = get_tesseract_command()

        self.assertTrue(command is None or isinstance(command, str))


if __name__ == "__main__":
    unittest.main()
